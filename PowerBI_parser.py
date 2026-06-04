# power.py — IES 90-Day Operational Report (Proofpoint)
import os
import numpy as np
from datetime import datetime, timedelta
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------- CONFIG ----------
CSV_PATH   = "alerts.csv"      # <-- your CSV filename
LOGO_PATH  = "td.png"
CHART_DIR  = "charts"
LOOKBACK   = 90

# EXACT columns from your file
COL_DATE     = "Alert Date"
COL_CREATED  = "Created"
COL_STATUS   = "Status"
COL_RULE     = "Alert Rule"
COL_SEVERITY = "Severity"
COL_ANALYST  = "Analyst Name"
COL_SIGNOFF  = "Sign-off status"

end   = datetime.now().date()
start = end - timedelta(days=LOOKBACK)
window_label = f"{start:%m/%d/%Y} – {end:%m/%d/%Y}"
OUTPUT_DOCX  = f"IES_90DayReport_{end:%Y-%m-%d}.docx"
os.makedirs(CHART_DIR, exist_ok=True)

# ---------- LOAD ----------
df = pd.read_csv(CSV_PATH)
df[COL_DATE] = pd.to_datetime(df[COL_DATE], errors="coerce")
df = df.dropna(subset=[COL_DATE])
df = df[(df[COL_DATE].dt.date >= start) & (df[COL_DATE].dt.date <= end)]

# ---------- METRICS ----------
alerts_received = len(df)
alerts_pending  = df[COL_STATUS].astype(str).str.lower().eq("pending").sum()
avg_per_day     = round(alerts_received / LOOKBACK, 1)

status_counts   = df[df[COL_STATUS].isin(["Non-Issue", "Issue"])][COL_STATUS].value_counts()
severity_counts = df[COL_SEVERITY].value_counts()

rule_counts = df[COL_RULE].value_counts()
top5  = rule_counts.head(5)
other = rule_counts.iloc[5:].sum()
rules_plot = pd.concat([top5, pd.Series({"Others": other})]) if other else top5

analyst_counts = (df[df[COL_STATUS].isin(["Non-Issue", "Issue"])]
                  [COL_ANALYST].value_counts())

# ---------- CHARTS ----------
TD_GREEN  = "#00B140"
TD_DARK   = "#2D2D2D"
TD_ACCENT = "#B0413E"
plt.rcParams.update({"font.family": "Calibri", "font.size": 11})

def style(ax):
    ax.spines[["top","right"]].set_visible(False)
    ax.grid(axis="y", linestyle="--", alpha=0.3)
    ax.set_axisbelow(True)

def save_bar(series, fname, ylabel, xlabel, color=TD_GREEN, figsize=(8,4.2)):
    if series.empty: return
    fig, ax = plt.subplots(figsize=figsize)
    bars = ax.bar(series.index.astype(str), series.values, color=color, width=0.6)
    ax.bar_label(bars, padding=3, fontsize=10)
    ax.set_ylabel(ylabel); ax.set_xlabel(xlabel)
    style(ax)
    plt.xticks(rotation=15, ha="right")
    plt.tight_layout(); plt.savefig(f"{CHART_DIR}/{fname}", dpi=160); plt.close()

def save_daily_trend(df, fname):
    # daily counts across the full 90-day window (fill missing days with 0)
    daily = (df.groupby(df[COL_DATE].dt.date).size()
               .reindex(pd.date_range(start, end).date, fill_value=0))
    fig, ax = plt.subplots(figsize=(11, 4.2))
    ax.bar(daily.index, daily.values, color=TD_GREEN, width=0.9, label="Daily alerts")

    # 7-day rolling avg as trend line
    trend = pd.Series(daily.values).rolling(7, min_periods=1).mean()
    ax.plot(daily.index, trend.values, color=TD_ACCENT, linewidth=2.2,
            label="7-day rolling avg")

    # linear trendline
    x = np.arange(len(daily))
    if len(x) > 1:
        m, b = np.polyfit(x, daily.values, 1)
        ax.plot(daily.index, m*x + b, color=TD_DARK, linestyle="--",
                linewidth=1.5, label="Linear trend")

    ax.set_ylabel("Alerts per day"); ax.set_xlabel("Date")
    style(ax)
    ax.legend(loc="upper left", frameon=False)
    fig.autofmt_xdate()
    plt.tight_layout(); plt.savefig(f"{CHART_DIR}/{fname}", dpi=160); plt.close()

# Build the charts
save_bar(severity_counts, "severity.png", "Number of Alerts", "Severity",
         color=TD_GREEN, figsize=(6.5,4.0))
save_bar(status_counts,   "status.png",   "Number of Alerts", "Disposition",
         color=TD_DARK, figsize=(6.5,4.0))
save_bar(rules_plot,      "rules.png",    "Volume of Alerts", "Alert Rule")
save_bar(analyst_counts,  "analysts.png", "Alerts Completed", "Analyst")
save_daily_trend(df, "daily_trend.png")

# ---------- DOC ----------
doc = Document()

hdr = doc.sections[0].header.paragraphs[0]
hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
if os.path.exists(LOGO_PATH):
    hdr.add_run().add_picture(LOGO_PATH, width=Inches(0.7))

def new_page(title):
    doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run("Insider Enhanced Surveillance")
    r.bold = True; r.font.size = Pt(22)
    doc.add_paragraph(f"90-Day Operational Report | {title}\n{window_label}")

# Cover
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\n\n\nInsider Enhanced Surveillance\n"
              "90-Day Operational Report – Proofpoint\n")
r.bold = True; r.font.size = Pt(28)
c = doc.add_paragraph(window_label); c.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Volume & Triage
new_page("Volume and Triage Metrics")

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Light Grid Accent 1"
hdr_cells = tbl.rows[0].cells
hdr_cells[0].text = "Alerts Received"
hdr_cells[1].text = "Alerts Pending"
hdr_cells[2].text = "Avg per Day"
row = tbl.add_row().cells
row[0].text = str(alerts_received)
row[1].text = str(int(alerts_pending))
row[2].text = str(avg_per_day)

doc.add_paragraph()  # spacer
doc.add_paragraph("Daily Alert Volume (90 days)").runs[0].bold = True
doc.add_picture(f"{CHART_DIR}/daily_trend.png", width=Inches(6.8))

doc.add_paragraph()
doc.add_paragraph("Severity Breakdown").runs[0].bold = True
doc.add_picture(f"{CHART_DIR}/severity.png", width=Inches(3.3))

doc.add_paragraph("Issue vs Non-Issue").runs[0].bold = True
doc.add_picture(f"{CHART_DIR}/status.png", width=Inches(3.3))

# Rules
new_page("Alerts by Rule Type Breakdown")
if os.path.exists(f"{CHART_DIR}/rules.png"):
    doc.add_picture(f"{CHART_DIR}/rules.png", width=Inches(6.5))

# Analyst
new_page("Alerts Completed by Analyst")
if os.path.exists(f"{CHART_DIR}/analysts.png"):
    doc.add_picture(f"{CHART_DIR}/analysts.png", width=Inches(6.5))

# Highlights
new_page("Program Highlights")
doc.add_paragraph("[ADD ACCOMPLISHMENTS HERE BEFORE EXPORTING TO PDF]")
doc.add_paragraph(f"There were {int(status_counts.get('Issue', 0))} "
                  f"identified issues over the 90-day period.")

doc.save(OUTPUT_DOCX)
print(f"Saved {OUTPUT_DOCX}")
